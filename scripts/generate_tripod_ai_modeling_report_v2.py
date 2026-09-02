#!/usr/bin/env python3
"""Create the TRIPOD+AI-oriented v2 report for the Timbio H. pylori run."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "artifacts" / "model_registry" / "model_results_timbio_hpylori_v2.csv"
REGISTRY = ROOT / "data" / "manifests" / "feature_registry_timbio_hpylori_v2.csv"
BOOTSTRAP = ROOT / "artifacts" / "metrics" / "bootstrap_optimism_timbio_hpylori_v2.csv"
SHAP = ROOT / "artifacts" / "metrics" / "shap_importance_timbio_hpylori_v2.csv"
DECISION = ROOT / "artifacts" / "metrics" / "decision_curve_timbio_hpylori_v2.csv"
FIGURES = ROOT / "artifacts" / "figures" / "timbio_hpylori_v2"
OUT_MD = ROOT / "docs" / "informe_modelado_tripod_ai_v2.md"
OUT_DOCX = ROOT / "docs" / "informe_modelado_tripod_ai_v2.docx"


def corrected_metrics(winner: pd.Series, bootstrap: pd.DataFrame) -> tuple[float, float]:
    return (
        float(winner["auc_oof"] - bootstrap["auc_optimism"].mean()),
        float(winner["brier_oof"] + bootstrap["brier_optimism"].mean()),
    )


def markdown_table(frame: pd.DataFrame, columns: list[str]) -> str:
    rounded = frame[columns].copy()
    for column in columns:
        if pd.api.types.is_numeric_dtype(rounded[column]):
            rounded[column] = rounded[column].map(lambda value: f"{value:.3f}")
    header = "| " + " | ".join(columns) + " |"
    divider = "|" + "|".join(["---"] * len(columns)) + "|"
    body = "\n".join("| " + " | ".join(str(value) for value in row) + " |" for row in rounded.itertuples(index=False, name=None))
    return "\n".join([header, divider, body])


def build_markdown(results: pd.DataFrame, registry: pd.DataFrame, bootstrap: pd.DataFrame, shap_table: pd.DataFrame, decision: pd.DataFrame) -> str:
    winner = results.iloc[0]
    corrected_auc, corrected_brier = corrected_metrics(winner, bootstrap)
    model_table = markdown_table(results, ["model_name", "auc_oof", "auc_oof_ci_low", "auc_oof_ci_high", "brier_oof", "calibration_slope_oof", "threshold_selected_oof", "sensitivity_oof", "specificity_oof"])
    exclusions = registry.loc[~registry["model_eligible"], ["variable_canonical", "reason"]]
    exclusion_list = "\n".join(f"- `{row.variable_canonical}`: {row.reason}" for row in exclusions.itertuples(index=False))
    shap_list = "\n".join(f"- `{row.encoded_feature}`: media |SHAP|={row.mean_abs_shap:.3f}" for row in shap_table.head(10).itertuples(index=False))
    model_decision = decision.loc[decision["strategy"].eq("model")]
    low_threshold = model_decision.iloc[0]
    high_threshold = model_decision.loc[model_decision["threshold"].eq(float(winner["threshold_selected_oof"]))].iloc[0]
    return f"""# Informe de modelado TRIPOD+AI v2: H. pylori en Timbío

Generado: {datetime.now(timezone.utc).isoformat(timespec="seconds")} UTC.

## Dictamen ejecutivo

Esta segunda iteración corrige los problemas metodológicos señalados en la revisión previa: fuente local sin publicar, entrada de modelado desidentificada, exclusión explícita de leakage, preprocesamiento dentro de los pliegues, validación cruzada anidada, bootstrap de optimismo, calibración y análisis de decisión. Aun con estas correcciones, no se observó discriminación clínica útil para el desenlace `Resultados_Helicobacter`. El mejor modelo fue `{winner.model_name}` con AUC fuera de pliegue {winner.auc_oof:.3f} (IC bootstrap {winner.auc_oof_ci_low:.3f}-{winner.auc_oof_ci_high:.3f}); el intervalo cruza 0.5. El desempeño corregido por optimismo fue AUC={corrected_auc:.3f}. Por tanto, el modelo no debe usarse para decisión clínica ni priorización de pacientes.

## 1. Pregunta y fuente de datos

Población: pacientes de Timbío de `biomarkers_timbio_with_sociodemographics_v2.xlsx`.

Outcome provisional: `Resultados_Helicobacter`, con `Positivo`=1 y `Negativo`=0. Se excluyeron los registros `No Registra`. La cohorte analítica final tuvo {int(winner.n_patients)} pacientes y {int(winner.n_positive)} positivos ({winner.prevalence:.1%}).

La elección de `Resultados_Helicobacter` responde a la indicación de esta iteración. Frente a `H_pylori_heces`, hubo 276 coincidencias informativas y un positivo adicional en `H_pylori_heces`; este detalle debe mantenerse documentado si el outcome cambia en una futura iteración.

## 2. Correcciones implementadas respecto a la revisión metodológica

- La fuente se trató como solo lectura y se creó una matriz de modelado sin identificadores directos en `data/interim/`.
- La política de predictores partió de las 63 variables obligatorias de la columna R del diccionario oficial.
- Se autorizaron {int(registry.model_eligible.sum())} predictores prediagnóstico. Se excluyeron datos sensibles, resultados, biomarcadores, endoscopia, histología, scores y riesgos derivados.
- Imputación, manejo de faltantes, agrupación de categorías raras, one-hot encoding y filtrado de varianza se ajustaron exclusivamente dentro de cada pliegue.
- Se compararon elastic net, random forest y gradient boosting bajo el mismo esquema de validación cruzada anidada.
- El stepwise no se usó como estrategia de selección.
- SHAP se calculó únicamente sobre el modelo final reconstruido bajo este pipeline corregido.

Variables obligatorias no elegibles:
{exclusion_list}

## 3. Validación y evaluación

Se usaron 5 pliegues externos estratificados y 3 pliegues internos para tuning por AUC. Las métricas principales se calcularon sobre predicciones fuera de pliegue. El modelo ganador se ajustó sobre la cohorte completa y se sometió a 200 réplicas bootstrap para estimar optimismo.

No se fijó el umbral en 0.5. Se exploraron umbrales entre 0.05 y 0.50; para obtener al menos 80% de sensibilidad, el umbral descriptivo del modelo ganador fue {winner.threshold_selected_oof:.2f}. Esto logró sensibilidad {winner.sensitivity_oof:.3f}, pero especificidad {winner.specificity_oof:.3f}, por lo que derivaría casi toda la población y no es operacionalmente útil.

## 4. Comparación de modelos

{model_table}

El random forest tuvo la mayor AUC fuera de pliegue, pero su magnitud fue mínima y compatible con azar. Gradient boosting tuvo menor Brier aparente fuera de pliegue, aunque su pendiente de calibración fue muy baja. Ningún modelo proporciona evidencia suficiente de utilidad predictiva.

## 5. Calibración, bootstrap y utilidad clínica

Para el modelo ganador: Brier fuera de pliegue={winner.brier_oof:.3f}; intercepto de calibración={winner.calibration_intercept_oof:.3f}; pendiente={winner.calibration_slope_oof:.3f}. La pendiente inferior a 1 indica sobreajuste. El bootstrap de optimismo estimó AUC corregida={corrected_auc:.3f} y Brier corregido={corrected_brier:.3f}.

En la curva de decisión, el beneficio neto del modelo fue igual a “tratar a todos” en umbrales 0.05-0.30. A 0.35 fue {model_decision.loc[model_decision.threshold.eq(0.35), 'net_benefit'].iloc[0]:.3f} y a {high_threshold.threshold:.2f} fue {high_threshold.net_benefit:.3f}; ambos no sostienen una estrategia de tamización selectiva. El umbral más bajo tuvo beneficio neto {low_threshold.net_benefit:.3f}, idéntico a tratar a todos.

## 6. Explicabilidad del modelo final

SHAP describe qué variables influyen en las probabilidades del random forest final, no evidencia de causalidad ni de desempeño útil. Las diez variables codificadas con mayor importancia global fueron:

{shap_list}

Estas señales son exploratorias. Dado que el AUC externo interno es cercano a 0.5, no deben convertirse en un score ni interpretarse como factores de riesgo confirmados.

## 7. Limitaciones y siguiente fase

- El tamaño de la cohorte es limitado: 89 eventos para 57 predictores base, antes de la expansión por categorías.
- La fuente integrada contiene biomarcadores y resultados posteriores, pero fueron excluidos deliberadamente de este modelo clínico-social base.
- La ausencia de señal puede reflejar que la infección activa no está suficientemente determinada por los predictores disponibles, errores de medición, heterogeneidad o selección de la población.
- Este análisis no reemplaza la futura fase de biomarcadores ni el desenlace histológico/endoscópico.

La siguiente iteración debe definir una pregunta separada para biomarcadores o para endoscopia/histología, con temporalidad clara, tamaño muestral suficiente y validación externa o temporal. No se recomienda optimizar más modelos sobre esta misma matriz para buscar una mejora marginal.

## 8. Reproducción y artefactos

- Configuración: `configs/timbio_hpylori_model_v2.yaml`
- Ejecución: `scripts/run_timbio_hpylori_model_v2.py`
- Registro de predictores: `data/manifests/feature_registry_timbio_hpylori_v2.csv`
- Matriz de entrada desidentificada: `data/interim/timbio_hpylori_model_input_deidentified_v2.parquet`
- Resultados, predicciones fuera de pliegue, bootstrap, calibración, curva de decisión, SHAP y figuras: `artifacts/` con sufijo `_v2`.
"""


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8)


def add_results_table(doc: Document, results: pd.DataFrame) -> None:
    columns = ["Modelo", "AUC (IC 95%)", "Brier", "Pendiente cal.", "Umbral", "Sens.", "Espec."]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, name in enumerate(columns):
        set_cell_text(table.rows[0].cells[idx], name, bold=True)
    for row in results.itertuples(index=False):
        values = [
            row.model_name,
            f"{row.auc_oof:.3f} ({row.auc_oof_ci_low:.3f}-{row.auc_oof_ci_high:.3f})",
            f"{row.brier_oof:.3f}",
            f"{row.calibration_slope_oof:.3f}",
            f"{row.threshold_selected_oof:.2f}",
            f"{row.sensitivity_oof:.3f}",
            f"{row.specificity_oof:.3f}",
        ]
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value)
    for row in table.rows:
        for idx, width in enumerate([1.25, 1.15, 0.65, 0.85, 0.55, 0.55, 0.55]):
            row.cells[idx].width = Inches(width)


def add_picture(doc: Document, path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.1))
    caption_p = doc.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.runs[0].italic = True
    caption_p.runs[0].font.size = Pt(8)


def build_docx(results: pd.DataFrame, registry: pd.DataFrame, bootstrap: pd.DataFrame, shap_table: pd.DataFrame, decision: pd.DataFrame) -> None:
    winner = results.iloc[0]
    corrected_auc, corrected_brier = corrected_metrics(winner, bootstrap)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(10)
    for style_name, size, color in [("Heading 1", 15, RGBColor(31, 79, 121)), ("Heading 2", 12, RGBColor(31, 79, 121))]:
        style = styles[style_name]; style.font.name = "Arial"; style.font.size = Pt(size); style.font.color.rgb = color

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Informe de modelado TRIPOD+AI v2")
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(22); run.font.color.rgb = RGBColor(31, 79, 121)
    subtitle = doc.add_paragraph("Modelo de factibilidad de H. pylori en pacientes de Timbío")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.name = "Arial"; subtitle.runs[0].font.size = Pt(13)
    generated = doc.add_paragraph(f"Generado: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.runs[0].font.size = Pt(9)

    doc.add_heading("Dictamen ejecutivo", level=1)
    p = doc.add_paragraph()
    p.add_run("Conclusión: ").bold = True
    p.add_run(f"tras corregir el pipeline y validar internamente, el mejor modelo alcanzó AUC fuera de pliegue {winner.auc_oof:.3f} (IC 95% {winner.auc_oof_ci_low:.3f}-{winner.auc_oof_ci_high:.3f}), compatible con discriminación no útil. No debe usarse para priorización clínica.")

    doc.add_heading("Población, outcome y gobernanza", level=1)
    doc.add_paragraph(f"La cohorte de Timbío incluyó {int(winner.n_patients)} pacientes con `Resultados_Helicobacter` válido: {int(winner.n_positive)} positivos ({winner.prevalence:.1%}) y {int(winner.n_patients - winner.n_positive)} negativos. Se excluyeron `No Registra`.")
    doc.add_paragraph(f"El diccionario oficial aportó 63 variables obligatorias de la columna R. El modelo utilizó {int(registry.model_eligible.sum())} predictores prediagnóstico autorizados. Se excluyeron identificadores, información sensible, biomarcadores, resultados, endoscopia, histología, scores, riesgos derivados, fechas y variables sin observaciones suficientes.")
    doc.add_paragraph("Las correcciones solicitadas en la revisión previa se aplicaron: copia de entrada desidentificada, preprocesamiento dentro del remuestreo, comparadores homogéneos, ausencia de stepwise como selección principal, bootstrap, calibración, análisis de decisión y SHAP sobre el modelo reconstruido.")

    doc.add_heading("Métodos de validación", level=1)
    doc.add_paragraph("Se compararon regresión logística elastic net, random forest y gradient boosting. El tuning se efectuó con validación cruzada anidada: cinco pliegues externos y tres internos, todos estratificados. La imputación, las categorías raras, el one-hot encoding y el filtrado de varianza se ajustaron solo con los datos de entrenamiento de cada pliegue.")
    doc.add_paragraph("El modelo ganador fue refit sobre la cohorte completa y evaluado mediante 200 réplicas bootstrap para estimar optimismo. Se reportan AUC con intervalo bootstrap, Brier, parámetros de calibración, umbrales no fijados en 0.5, curva de decisión y explicación SHAP.")

    doc.add_heading("Resultados comparativos", level=1)
    add_results_table(doc, results)
    doc.add_paragraph(f"El ganador por AUC fue random forest, pero la diferencia frente a los otros modelos fue mínima. Su AUC corregida por optimismo fue {corrected_auc:.3f}; Brier corregido {corrected_brier:.3f}. La pendiente de calibración de {winner.calibration_slope_oof:.3f} indica sobreajuste.")
    add_picture(doc, FIGURES / "auc_comparison_v2.png", "Figura 1. AUC fuera de pliegue con intervalos bootstrap.")
    add_picture(doc, FIGURES / "roc_oof_v2.png", "Figura 2. Curvas ROC calculadas con predicciones fuera de pliegue.")

    doc.add_heading("Calibración y utilidad clínica", level=1)
    doc.add_paragraph(f"El umbral descriptivo de {winner.threshold_selected_oof:.2f} se eligió para alcanzar al menos 80% de sensibilidad. Alcanzó sensibilidad {winner.sensitivity_oof:.3f}, pero especificidad {winner.specificity_oof:.3f}; en la práctica derivaría casi todos los pacientes.")
    doc.add_paragraph("La curva de decisión no muestra una ventaja clínicamente útil frente a tratar a todos en los umbrales bajos. A los umbrales donde deja de coincidir, el beneficio neto es negativo o no supera tratar a ninguno. No se justifica un umbral operativo.")
    add_picture(doc, FIGURES / "calibration_oof_v2.png", "Figura 3. Curvas de calibración fuera de pliegue.")
    add_picture(doc, FIGURES / "decision_curve_v2.png", "Figura 4. Curva de decisión del modelo ganador.")

    doc.add_heading("Explicabilidad", level=1)
    doc.add_paragraph("SHAP se presenta como descripción del modelo final, no como evidencia de causalidad ni como criterio de adopción. Las señales principales se concentran en IMC, peso, edad, características de vivienda y variables de estilo de vida, pero deben interpretarse con extrema cautela dada la discriminación global débil.")
    add_picture(doc, FIGURES / "shap_importance_v2.png", "Figura 5. Importancia global SHAP del modelo final.")

    doc.add_heading("Limitaciones y recomendaciones", level=1)
    for text in [
        "La cohorte tiene 89 eventos para 57 predictores base; la expansión por categorías limita la estabilidad.",
        "El outcome es provisional y debe mantenerse trazable respecto a H_pylori_heces y futuros resultados clínicos.",
        "La ausencia de señal no debe resolverse probando más algoritmos sobre esta misma matriz; sería una fuente adicional de optimismo.",
        "La siguiente fase debe separar una pregunta de biomarcadores o de endoscopia/histología, con temporalidad explícita, mayor muestra y validación externa o temporal.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Reproducción", level=1)
    doc.add_paragraph("La ejecución se reproduce con `configs/timbio_hpylori_model_v2.yaml` y `scripts/run_timbio_hpylori_model_v2.py` usando `.venv-model-v2`. Los archivos derivados se almacenan localmente con sufijo `_v2`; las fuentes originales no se modificaron.")
    doc.save(OUT_DOCX)


def main() -> None:
    results = pd.read_csv(RESULTS).sort_values("auc_oof", ascending=False)
    registry = pd.read_csv(REGISTRY)
    bootstrap = pd.read_csv(BOOTSTRAP)
    shap_table = pd.read_csv(SHAP)
    decision = pd.read_csv(DECISION)
    OUT_MD.write_text(build_markdown(results, registry, bootstrap, shap_table, decision))
    build_docx(results, registry, bootstrap, shap_table, decision)
    print(f"Markdown: {OUT_MD}")
    print(f"DOCX: {OUT_DOCX}")


if __name__ == "__main__":
    main()
